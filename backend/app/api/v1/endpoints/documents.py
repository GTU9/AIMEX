"""
DOCUMENTS 테이블 API 엔드포인트
"""

from typing import List, Optional, Dict, Any
from fastapi import APIRouter, Depends, HTTPException, Query, File, UploadFile, Form
from pydantic import BaseModel
from sqlalchemy.orm import Session
from sqlalchemy import select
import logging
import tempfile
import os
import re
from datetime import datetime

from app.database import get_db
from app.models.rag import Documents
from app.services.rag_document_service import get_rag_document_service

logger = logging.getLogger(__name__)

router = APIRouter()


# ==================== Response Models ====================


class DocumentResponse(BaseModel):
    """문서 응답 모델"""

    documents_id: str
    documents_name: str
    file_size: Optional[int]
    file_path: str
    is_vectorized: int
    created_at: Optional[str]


class DocumentListResponse(BaseModel):
    """문서 목록 응답 모델"""

    documents: List[DocumentResponse]
    total_count: int


class DocumentUploadResponse(BaseModel):
    """문서 업로드 응답 모델"""

    success: bool
    message: str
    documents_id: Optional[str] = None
    file_path: Optional[str] = None
    file_size: Optional[int] = None


class VectorizationUpdateResponse(BaseModel):
    """벡터화 상태 업데이트 응답 모델"""

    success: bool
    message: str
    documents_id: str
    is_vectorized: int


class S3FileInfo(BaseModel):
    """S3 파일 정보 모델"""

    key: str
    filename: str
    size: int
    last_modified: str
    presigned_url: Optional[str] = None


class S3FileListResponse(BaseModel):
    """S3 파일 목록 응답 모델"""

    files: List[S3FileInfo]
    total_count: int


# ==================== Document Endpoints ====================


@router.post("/upload", response_model=DocumentUploadResponse)
async def upload_document(
    file: UploadFile = File(..., description="PDF 또는 텍스트 파일"),
    influencer_id: Optional[str] = Form(None, description="소유 인플루언서 ID"),
    db: Session = Depends(get_db),
):
    """문서를 로컬 파일시스템에 저장하고 DB에 기록 (S3 제거, NAS 로컬 전환)."""
    import uuid
    from app.core.config import settings

    try:
        logger.info(f"📥 문서 업로드 시작: {file.filename} (influencer_id={influencer_id})")

        fname = (file.filename or "").lower()
        if not (fname.endswith(".pdf") or fname.endswith(".txt") or fname.endswith(".md") or fname.endswith(".docx")):
            raise HTTPException(status_code=400, detail="PDF/DOCX/TXT/MD 파일만 업로드 가능합니다.")

        content = await file.read()
        if len(content) > 10 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="파일 크기는 10MB를 초과할 수 없습니다.")

        # 로컬 저장 (LOCAL_STORAGE_PATH 의 형제 디렉터리 documents/)
        base_dir = os.path.join(
            os.path.dirname(settings.LOCAL_STORAGE_PATH.rstrip("/")) or ".", "documents"
        )
        os.makedirs(base_dir, exist_ok=True)
        documents_id = str(uuid.uuid4())
        save_path = os.path.join(base_dir, f"{documents_id}_{file.filename}")
        with open(save_path, "wb") as f:
            f.write(content)
        logger.info(f"✅ 로컬 저장: {save_path} ({len(content)} bytes)")

        row = Documents(
            documents_id=documents_id,
            documents_name=file.filename,
            file_size=len(content),
            file_path=save_path,
            is_vectorized=0,
            influencer_id=influencer_id,
        )
        db.add(row)
        db.commit()

        return DocumentUploadResponse(
            success=True,
            message="문서가 로컬에 저장되었습니다.",
            documents_id=documents_id,
            file_path=save_path,
            file_size=len(content),
        )

    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        logger.error(f"❌ 문서 업로드 실패: {e}")
        raise HTTPException(status_code=500, detail=f"문서 업로드 실패: {str(e)}")


def _chunk_plaintext(raw: str, max_len: int = 400) -> list:
    """txt/md/docx 텍스트를 마크다운 헤더(섹션)·줄 단위로 의미 있는 청크로 분할.

    rag_service._preprocess_text 는 개행을 모두 공백으로 합쳐 문서를 한 덩어리로 만들어
    검색 정밀도가 떨어진다. 여기서는 섹션/항목 구조를 보존해 작은 청크로 나눠 RAG 적중률을 높인다.
    """
    raw = (raw or "").replace("\r\n", "\n")
    sections = re.split(r"\n(?=#{1,6}\s)", raw)  # 마크다운 헤더 기준 섹션 분리
    chunks = []
    for sec in sections:
        header = ""
        buf = ""
        for ln in sec.split("\n"):
            ln = ln.strip()
            if not ln:
                continue
            if re.match(r"^#{1,6}\s", ln):
                header = ln.lstrip("#").strip()
                continue
            ln = ln.lstrip("-*•").strip()
            if not ln:
                continue
            piece = f"[{header}] {ln}" if header else ln
            if len(buf) + len(piece) > max_len:
                if buf:
                    chunks.append(buf.strip())
                buf = piece
            else:
                buf = (buf + " " + piece).strip()
        if buf:
            chunks.append(buf.strip())
    return [c for c in chunks if len(c.strip()) >= 5]


async def _extract_chunks(path: str, proc) -> list:
    """파일 형식별 텍스트 추출 + 청킹 (pdf/docx/txt/md)."""
    low = path.lower()
    if low.endswith(".pdf"):
        qa = await proc.process_pdf(path)
        return [q["answer"] for q in qa if q.get("answer")]
    if low.endswith(".docx"):
        from docx import Document as DocxDocument
        d = DocxDocument(path)
        parts = [p.text for p in d.paragraphs if p.text and p.text.strip()]
        for tbl in d.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return _chunk_plaintext("\n".join(parts))
    # txt / md
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return _chunk_plaintext(f.read())


@router.post("/{documents_id}/vectorize")
async def vectorize_document(documents_id: str, db: Session = Depends(get_db)):
    """문서를 청킹→임베딩(Modal)→Chroma 저장하여 RAG 검색 가능 상태로 만든다."""
    from app.services.rag_service import RAGDocumentProcessor, RAGConfig
    from app.services.embedding_client import embed_texts
    from app.services.rag_vector_store import get_vector_store

    doc = db.query(Documents).filter(Documents.documents_id == documents_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="문서를 찾을 수 없습니다.")
    if not doc.influencer_id:
        raise HTTPException(status_code=400, detail="문서에 influencer_id 가 없습니다.")

    proc = RAGDocumentProcessor(RAGConfig())
    path = doc.file_path
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail="원본 파일이 존재하지 않습니다.")

    chunks = await _extract_chunks(path, proc)

    chunks = [c for c in chunks if c and c.strip()]
    if not chunks:
        raise HTTPException(status_code=400, detail="추출된 텍스트가 없습니다.")

    vecs = await embed_texts(chunks)
    store = get_vector_store()
    store.delete_by_influencer(doc.influencer_id)  # 재벡터화 시 해당 인플루언서 기존 벡터 제거
    store.upsert([
        {"text": c, "embedding": vecs[i], "influencer_id": doc.influencer_id,
         "source": doc.documents_name, "chunk_id": i}
        for i, c in enumerate(chunks)
    ])
    doc.is_vectorized = 1
    db.commit()
    logger.info(f"✅ 벡터화 완료: {documents_id} chunks={len(chunks)} influencer={doc.influencer_id}")
    return {"documents_id": documents_id, "chunks": len(chunks), "is_vectorized": 1}


async def _rebuild_influencer_vectors(influencer_id: str, db: Session) -> dict:
    """해당 인플루언서의 현재 저장된 모든 문서로 Chroma 벡터를 재구축한다.

    매 호출마다 인플루언서 벡터를 전부 비우고 현재 문서 전체로 다시 쌓는다.
    문서가 없으면(또는 텍스트가 없으면) 기존 벡터만 삭제 → 고아 벡터가 남지 않는다.
    (업로드/삭제/재반영 모두 이 함수로 수렴시켜 항상 최신 상태 유지)
    """
    from app.services.rag_service import RAGDocumentProcessor, RAGConfig
    from app.services.embedding_client import embed_texts
    from app.services.rag_vector_store import get_vector_store

    store = get_vector_store()
    docs = db.query(Documents).filter(Documents.influencer_id == influencer_id).all()

    if not docs:
        store.delete_by_influencer(influencer_id)
        logger.info(f"🧹 인플루언서 {influencer_id} 문서 없음 → 벡터 전체 삭제")
        return {"influencer_id": influencer_id, "documents": 0,
                "embedded_documents": 0, "chunks": 0, "skipped": []}

    proc = RAGDocumentProcessor(RAGConfig())
    rows: list[dict] = []   # {text, source}
    skipped: list[str] = []

    for doc in docs:
        path = doc.file_path
        if not path or not os.path.exists(path):
            skipped.append(doc.documents_name)
            continue
        try:
            chunks = await _extract_chunks(path, proc)
        except Exception as e:
            logger.warning(f"⚠️ 문서 처리 실패 {doc.documents_name}: {e}")
            skipped.append(doc.documents_name)
            continue
        for c in chunks:
            if c and c.strip():
                rows.append({"text": c.strip(), "source": doc.documents_name})

    skipped_set = set(skipped)

    if not rows:
        # 임베딩할 텍스트 없음 → 벡터 비우고 미반영 처리
        store.delete_by_influencer(influencer_id)
        for doc in docs:
            doc.is_vectorized = 0
        db.commit()
        return {"influencer_id": influencer_id, "documents": len(docs),
                "embedded_documents": 0, "chunks": 0, "skipped": skipped}

    texts = [r["text"] for r in rows]
    vecs = await embed_texts(texts)

    store.delete_by_influencer(influencer_id)  # 전체 재구축
    store.upsert([
        {
            "text": texts[i],
            "embedding": vecs[i],
            "influencer_id": influencer_id,
            "source": rows[i]["source"],
            "chunk_id": i,
        }
        for i in range(len(texts))
    ])

    for doc in docs:
        doc.is_vectorized = 0 if doc.documents_name in skipped_set else 1
    db.commit()

    embedded_docs = len(docs) - len(skipped)
    logger.info(
        f"✅ 인플루언서 {influencer_id} 벡터 재구축: 문서 {embedded_docs}/{len(docs)}개, 청크 {len(texts)}개"
    )
    return {
        "influencer_id": influencer_id,
        "documents": len(docs),
        "embedded_documents": embedded_docs,
        "chunks": len(texts),
        "skipped": skipped,
    }


@router.post("/by-influencer/{influencer_id}/vectorize")
async def vectorize_influencer_documents(influencer_id: str, db: Session = Depends(get_db)):
    """해당 인플루언서의 저장된 '모든' 문서를 청킹→임베딩→Chroma에 일괄 재구축한다.

    업로드는 저장만 하고, 임베딩은 이 버튼(엔드포인트)으로 한 번에 수행한다.
    """
    docs_exist = (
        db.query(Documents).filter(Documents.influencer_id == influencer_id).first()
    )
    if not docs_exist:
        raise HTTPException(status_code=404, detail="해당 인플루언서의 문서가 없습니다.")

    result = await _rebuild_influencer_vectors(influencer_id, db)
    if result["chunks"] == 0:
        raise HTTPException(status_code=400, detail="임베딩할 텍스트가 없습니다.")
    return result


@router.get("/by-influencer/{influencer_id}", response_model=DocumentListResponse)
async def list_documents_by_influencer(
    influencer_id: str, db: Session = Depends(get_db)
):
    """특정 인플루언서에 속한 문서 목록 조회."""
    rows = (
        db.query(Documents)
        .filter(Documents.influencer_id == influencer_id)
        .order_by(Documents.created_at.desc())
        .all()
    )
    items = [
        DocumentResponse(
            documents_id=r.documents_id,
            documents_name=r.documents_name,
            file_size=r.file_size,
            file_path=r.file_path,
            is_vectorized=r.is_vectorized or 0,
            created_at=r.created_at.isoformat() if r.created_at else None,
        )
        for r in rows
    ]
    return DocumentListResponse(documents=items, total_count=len(items))


@router.put("/{documents_id}/vectorization", response_model=VectorizationUpdateResponse)
async def update_vectorization_status(
    documents_id: str,
    is_vectorized: int = Form(1, description="벡터화 상태 (1: 완료, 0: 미완료)"),
    db: Session = Depends(get_db),
):
    """문서의 벡터화 상태 업데이트"""
    try:
        rag_document_service = get_rag_document_service()

        success = rag_document_service.update_vectorization_status(
            documents_id=documents_id, db=db, is_vectorized=is_vectorized
        )

        if success:
            status_text = "완료" if is_vectorized == 1 else "미완료"
            return VectorizationUpdateResponse(
                success=True,
                message=f"벡터화 상태가 '{status_text}'로 업데이트되었습니다.",
                documents_id=documents_id,
                is_vectorized=is_vectorized,
            )
        else:
            raise HTTPException(status_code=404, detail="문서를 찾을 수 없습니다.")

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"벡터화 상태 업데이트 실패: {e}")
        raise HTTPException(
            status_code=500, detail=f"벡터화 상태 업데이트 실패: {str(e)}"
        )


@router.put("/reset-vectorization", response_model=VectorizationUpdateResponse)
async def reset_all_vectorization_status(db: Session = Depends(get_db)):
    """모든 문서의 벡터화 상태를 0으로 초기화"""
    try:
        rag_document_service = get_rag_document_service()

        success = rag_document_service.reset_all_vectorization_status(db=db)

        if success:
            return VectorizationUpdateResponse(
                success=True,
                message="모든 문서의 벡터화 상태가 초기화되었습니다.",
                documents_id="all",
                is_vectorized=0,
            )
        else:
            raise HTTPException(
                status_code=500, detail="벡터화 상태 초기화에 실패했습니다."
            )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"벡터화 상태 초기화 실패: {e}")
        raise HTTPException(
            status_code=500, detail=f"벡터화 상태 초기화 실패: {str(e)}"
        )


@router.get("/vectorized", response_model=DocumentListResponse)
async def get_vectorized_documents(
    limit: int = Query(50, description="조회 제한 수"),
    offset: int = Query(0, description="조회 시작 위치"),
    db: Session = Depends(get_db),
):
    """벡터화된 문서 목록 조회"""
    try:
        rag_document_service = get_rag_document_service()

        # 벡터화된 문서만 조회 (is_vectorized = 1)
        documents = await rag_document_service.get_vectorized_documents(
            db=db, limit=limit, offset=offset
        )

        return DocumentListResponse(documents=documents, total_count=len(documents))

    except Exception as e:
        logger.error(f"벡터화된 문서 목록 조회 실패: {e}")
        raise HTTPException(
            status_code=500, detail=f"벡터화된 문서 목록 조회 실패: {str(e)}"
        )


@router.get("/{documents_id}/download")
async def download_document(documents_id: str, db: Session = Depends(get_db)):
    """문서 다운로드 (Presigned URL 반환)"""
    try:
        rag_document_service = get_rag_document_service()
        document = await rag_document_service.get_document_by_id(
            documents_id=documents_id, db=db
        )

        if not document:
            raise HTTPException(status_code=404, detail="문서를 찾을 수 없습니다.")

        # S3에서 presigned URL 생성
        try:
            from app.services.s3_service import get_s3_service

            s3_service = get_s3_service()

            if s3_service.is_available():
                # S3 키 추출 (URL에서 키 부분만)
                file_path = document["file_path"]
                s3_key = file_path.split(".com/")[-1] if ".com/" in file_path else file_path

                # Presigned URL 생성 (24시간 유효)
                presigned_url = s3_service.generate_presigned_url(
                    s3_key, expiration=86400
                )

                if presigned_url:
                    return {
                        "success": True,
                        "download_url": presigned_url,
                        "filename": document["documents_name"],
                        "expires_in": 86400,
                    }
                else:
                    raise HTTPException(
                        status_code=500, detail="다운로드 URL 생성에 실패했습니다."
                    )
            else:
                raise HTTPException(
                    status_code=500, detail="S3 서비스를 사용할 수 없습니다."
                )

        except Exception as s3_error:
            logger.error(f"S3 다운로드 URL 생성 실패: {s3_error}")
            raise HTTPException(
                status_code=500, detail=f"다운로드 URL 생성 실패: {str(s3_error)}"
            )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"문서 다운로드 실패: {e}")
        raise HTTPException(status_code=500, detail=f"문서 다운로드 실패: {str(e)}")


@router.get("/s3", response_model=S3FileListResponse)
async def list_s3_documents(
    prefix: str = Query("documents/", description="S3 파일 경로 접두사"),
    include_presigned: bool = Query(True, description="Presigned URL 포함 여부"),
    db: Session = Depends(get_db),
):
    """S3에서 문서 파일 목록 조회"""
    try:
        from app.services.s3_service import get_s3_service

        s3_service = get_s3_service()

        if not s3_service.is_available():
            raise HTTPException(
                status_code=500, detail="S3 서비스를 사용할 수 없습니다."
            )

        # S3에서 파일 목록 조회
        if include_presigned:
            s3_files = s3_service.list_files_with_presigned_urls(prefix=prefix)
        else:
            s3_keys = s3_service.list_files(prefix=prefix)
            s3_files = []
            for key in s3_keys:
                # 파일명 추출
                filename = key.split("/")[-1]
                # 기본 정보만 포함
                s3_files.append(
                    {
                        "key": key,
                        "filename": filename,
                        "size": 0,  # 크기 정보는 별도 조회 필요
                        "last_modified": "",
                        "presigned_url": None,
                    }
                )

        # PDF 파일만 필터링
        pdf_files = [
            file for file in s3_files if file["filename"].lower().endswith(".pdf")
        ]

        return S3FileListResponse(files=pdf_files, total_count=len(pdf_files))

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"S3 문서 목록 조회 실패: {e}")
        raise HTTPException(status_code=500, detail=f"S3 문서 목록 조회 실패: {str(e)}")


@router.get("", response_model=DocumentListResponse)
async def get_documents(
    limit: int = Query(50, description="조회 제한 수"),
    offset: int = Query(0, description="조회 시작 위치"),
    db: Session = Depends(get_db),
):
    """문서 목록 조회"""
    try:
        rag_document_service = get_rag_document_service()

        # All documents lookup
        documents = await rag_document_service.get_all_documents(
            db=db, limit=limit, offset=offset
        )

        return DocumentListResponse(documents=documents, total_count=len(documents))

    except Exception as e:
        logger.error(f"문서 목록 조회 실패: {e}")
        raise HTTPException(status_code=500, detail=f"문서 목록 조회 실패: {str(e)}")


@router.get("/stats")
async def get_document_stats(db: Session = Depends(get_db)):
    """문서 통계 조회

    주의: 정적 경로(/stats)는 동적 경로(/{documents_id})보다 먼저 선언해야
    'stats'가 문서 ID로 매칭되는 것을 방지한다.
    """
    try:
        rag_document_service = get_rag_document_service()
        stats = await rag_document_service.get_all_document_stats(db=db)
        return stats

    except Exception as e:
        logger.error(f"문서 통계 조회 실패: {e}")
        raise HTTPException(status_code=500, detail=f"문서 통계 조회 실패: {str(e)}")


@router.get("/{documents_id}", response_model=DocumentResponse)
async def get_document(documents_id: str, db: Session = Depends(get_db)):
    """특정 문서 조회"""
    try:
        rag_document_service = get_rag_document_service()
        document = await rag_document_service.get_document_by_id(
            documents_id=documents_id, db=db
        )

        if not document:
            raise HTTPException(status_code=404, detail="문서를 찾을 수 없습니다.")

        return DocumentResponse(**document)

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"문서 조회 실패: {e}")
        raise HTTPException(status_code=500, detail=f"문서 조회 실패: {str(e)}")


@router.delete("/{documents_id}")
async def delete_document(
    documents_id: str,
    delete_from_s3: bool = Query(False, description="S3에서도 삭제할지 여부"),
    db: Session = Depends(get_db),
):
    """문서 삭제 (삭제 후 남은 문서로 벡터 재구축 → 고아 벡터 제거)"""
    try:
        # 삭제 전 소속 인플루언서/벡터화 여부 확보
        doc = db.query(Documents).filter(Documents.documents_id == documents_id).first()
        influencer_id = doc.influencer_id if doc else None
        was_vectorized = bool(doc.is_vectorized) if doc else False

        rag_document_service = get_rag_document_service()
        success = await rag_document_service.delete_document(
            documents_id=documents_id, db=db, delete_from_s3=delete_from_s3
        )

        if not success:
            raise HTTPException(status_code=404, detail="문서를 찾을 수 없습니다.")

        # 벡터화됐던 문서라면 Chroma 벡터를 남은 문서 기준으로 동기화 (best-effort)
        rebuilt = None
        if influencer_id and was_vectorized:
            try:
                rebuilt = await _rebuild_influencer_vectors(influencer_id, db)
                logger.info(f"🔄 문서 삭제 후 벡터 동기화: {rebuilt}")
            except Exception as e:
                logger.warning(f"⚠️ 삭제 후 벡터 동기화 실패(계속): {e}")

        return {"message": "문서가 성공적으로 삭제되었습니다.", "vectors_resynced": rebuilt is not None}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"문서 삭제 실패: {e}")
        raise HTTPException(status_code=500, detail=f"문서 삭제 실패: {str(e)}")
